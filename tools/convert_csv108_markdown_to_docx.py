from __future__ import annotations

import argparse
import sys
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


AUTHOR = "CSV108 Markdown DOCX Converter"
DOC_CREATED = datetime(2026, 6, 19, 9, 0, 0)
CONTENT_WIDTH_DXA = 9360
LANDSCAPE_CONTENT_WIDTH_DXA = 12960
TABLE_INDENT_DXA = 120
DXA_PER_INCH = 1440

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 89, 89)
BLACK = RGBColor(0, 0, 0)
WHITE = RGBColor(255, 255, 255)
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
NOTE_FILL = "F7F9FC"
CONTROL_FILL = "D9EAF7"

REQUIRED_META = {
    "document_number",
    "title",
    "project",
    "system",
    "lifecycle_phase",
    "version",
    "effective_or_record_date",
    "case_status",
    "prepared_by",
    "reviewed_by",
    "approved_by",
}


class MarkdownTableError(ValueError):
    pass


@dataclass(frozen=True)
class MarkdownBlock:
    kind: str
    text: str = ""
    level: int = 0
    rows: list[list[str]] = field(default_factory=list)
    start_line: int = 0
    language: str = ""


def parse_front_matter(text: str) -> tuple[dict[str, str], str]:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    if not normalized.startswith("---\n"):
        raise ValueError("Markdown file must start with YAML-like front matter.")
    parts = normalized.split("---\n", 2)
    if len(parts) != 3:
        raise ValueError("Markdown front matter must be closed with ---.")

    front = parts[1]
    body = parts[2].lstrip("\n")
    meta: dict[str, str] = {}
    for line in front.splitlines():
        if not line.strip():
            continue
        if ":" not in line:
            raise ValueError(f"Invalid front matter line: {line}")
        key, value = line.split(":", 1)
        meta[key.strip()] = value.strip().strip('"').strip("'")
    return meta, body


def validate_metadata(meta: dict[str, str], source_name: str) -> None:
    missing = sorted(REQUIRED_META - set(meta))
    if missing:
        raise ValueError(f"{source_name}: missing required front matter keys: {', '.join(missing)}")


def apply_known_markdown_repairs(path: Path, text: str) -> str:
    name = path.name
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    if name == "012_Data_Integrity_Risk_Assessment.md":
        for idx, line in enumerate(lines):
            if line.startswith("| DIRA-01 |") and line.count("|") - 1 == 8:
                cells = _parse_pipe_row(line)
                cells = cells[:4] + [f"{cells[4]}; {cells[5]}"] + cells[6:]
                lines[idx] = _format_pipe_row(cells)
    if name == "076_System_Handover_and_Support_Model.md":
        for idx, line in enumerate(lines):
            if line.startswith("| Quality oversight |") and line.count("|") - 1 == 4:
                cells = _parse_pipe_row(line)
                cells.append("On-call escalation")
                lines[idx] = _format_pipe_row(cells)
    return "\n".join(lines)


def parse_markdown_blocks(body: str, *, source_name: str) -> list[MarkdownBlock]:
    lines = body.splitlines()
    blocks: list[MarkdownBlock] = []
    idx = 0
    paragraph: list[str] = []
    paragraph_start = 0

    def flush_paragraph() -> None:
        nonlocal paragraph, paragraph_start
        if paragraph:
            blocks.append(MarkdownBlock(kind="paragraph", text=" ".join(part.strip() for part in paragraph), start_line=paragraph_start))
            paragraph = []
            paragraph_start = 0

    while idx < len(lines):
        raw = lines[idx]
        line = raw.rstrip()
        line_no = idx + 1

        if not line.strip():
            flush_paragraph()
            idx += 1
            continue

        if line.startswith("```"):
            flush_paragraph()
            language = line[3:].strip()
            code_start = line_no
            code_lines: list[str] = []
            idx += 1
            while idx < len(lines) and not lines[idx].startswith("```"):
                code_lines.append(lines[idx])
                idx += 1
            if idx < len(lines):
                idx += 1
            blocks.append(MarkdownBlock(kind="code", text="\n".join(code_lines), language=language, start_line=code_start))
            continue

        if line.startswith("# "):
            flush_paragraph()
            blocks.append(MarkdownBlock(kind="heading", text=line[2:].strip(), level=1, start_line=line_no))
            idx += 1
            continue
        if line.startswith("## "):
            flush_paragraph()
            blocks.append(MarkdownBlock(kind="heading", text=line[3:].strip(), level=2, start_line=line_no))
            idx += 1
            continue
        if line.startswith("### "):
            flush_paragraph()
            blocks.append(MarkdownBlock(kind="heading", text=line[4:].strip(), level=3, start_line=line_no))
            idx += 1
            continue
        if line.startswith(">"):
            flush_paragraph()
            quote_lines: list[str] = []
            quote_start = line_no
            while idx < len(lines) and lines[idx].startswith(">"):
                quote_lines.append(lines[idx].lstrip(">").strip())
                idx += 1
            blocks.append(MarkdownBlock(kind="quote", text=" ".join(quote_lines), start_line=quote_start))
            continue
        if line.startswith("- "):
            flush_paragraph()
            blocks.append(MarkdownBlock(kind="bullet", text=line[2:].strip(), start_line=line_no))
            idx += 1
            continue
        if line.strip().startswith("|") and "|" in line:
            flush_paragraph()
            table_lines: list[tuple[int, str]] = []
            table_start = line_no
            while idx < len(lines) and lines[idx].strip().startswith("|") and "|" in lines[idx]:
                table_lines.append((idx + 1, lines[idx].rstrip()))
                idx += 1
            blocks.append(MarkdownBlock(kind="table", rows=_parse_table(table_lines, source_name=source_name, start_line=table_start), start_line=table_start))
            continue

        if not paragraph:
            paragraph_start = line_no
        paragraph.append(line)
        idx += 1

    flush_paragraph()
    return blocks


def _parse_pipe_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _format_pipe_row(cells: list[str]) -> str:
    return "| " + " | ".join(cells) + " |"


def _is_separator_row(cells: list[str]) -> bool:
    if not cells:
        return False
    for cell in cells:
        stripped = cell.strip()
        if len(stripped) < 3:
            return False
        if set(stripped) - {"-", ":"}:
            return False
    return True


def _parse_table(table_lines: list[tuple[int, str]], *, source_name: str, start_line: int) -> list[list[str]]:
    parsed = [_parse_pipe_row(line) for _, line in table_lines]
    counts = [len(row) for row in parsed]
    if len(set(counts)) != 1:
        raise MarkdownTableError(
            f"{source_name}: Markdown table starting at line {start_line} has inconsistent column counts: {counts}"
        )
    if len(parsed) > 1 and _is_separator_row(parsed[1]):
        parsed.pop(1)
    return parsed


def set_run_font(
    run,
    *,
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    font_name: str = "Calibri",
):
    run.font.name = font_name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    _configure_section(section, landscape=False)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def _configure_section(section, *, landscape: bool) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def set_core_properties(doc: Document, meta: dict[str, str]) -> None:
    props = doc.core_properties
    props.author = AUTHOR
    props.last_modified_by = AUTHOR
    props.title = meta["title"]
    props.subject = "GMP CSV controlled-document conversion"
    props.category = "GMP CSV"
    props.keywords = "CSV, GMP, validation, DOCX, controlled document"
    props.comments = "Converted from CSV108 Markdown source. Requires site-specific QA review before operational use."
    props.created = DOC_CREATED
    props.modified = DOC_CREATED


def add_page_field(paragraph, instruction: str) -> None:
    run_begin = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_char_begin)

    run_instr = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" {instruction} "
    run_instr._r.append(instr_text)

    run_sep = paragraph.add_run()
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    run_sep._r.append(fld_char_sep)

    run_result = paragraph.add_run("1")
    set_run_font(run_result, size=9, color=MUTED)

    run_end = paragraph.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_char_end)


def add_header_footer(doc: Document, meta: dict[str, str]) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = ""
    left = header.add_run(f"{meta.get('project_code', 'QCL')} | {meta['system']}")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    header.add_run("\t")
    right = header.add_run(meta["document_number"])
    set_run_font(right, size=8.5, color=MUTED)
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer.paragraphs[0]
    footer.text = ""
    run = footer.add_run(f"{meta['document_number']} | Version {meta['version']} | {meta['case_status']} | Page ")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_field(footer, "PAGE")
    run = footer.add_run(" of ")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_field(footer, "NUMPAGES")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_title_block(doc: Document, meta: dict[str, str]) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(2)
    run = kicker.add_run("GMP CSV Controlled Document")
    set_run_font(run, size=9.5, color=MUTED, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run(meta["title"])
    set_run_font(run, size=22, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run(meta["project"])
    set_run_font(run, size=11.5, color=MUTED)

    add_metadata_table(
        doc,
        [
            ("Document number", meta["document_number"]),
            ("System", meta["system"]),
            ("Lifecycle phase", meta["lifecycle_phase"]),
            ("Version", meta["version"]),
            ("Effective/record date", meta["effective_or_record_date"]),
            ("Status", meta["case_status"]),
        ],
        header=("Control Field", "Value"),
        fill=CONTROL_FILL,
    )


def add_metadata_table(doc: Document, rows: list[tuple[str, str]], *, header: tuple[str, str], fill: str = HEADER_FILL) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, CONTENT_WIDTH_DXA - 2700], table_width_dxa=CONTENT_WIDTH_DXA)
    set_row_repeat_header(table.rows[0])
    for idx, value in enumerate(header):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, fill)
        write_cell(cell, value, bold=True, font_size=9.5)
    for label, value in rows:
        row = table.add_row()
        write_cell(row.cells[0], label, bold=True, font_size=9.5)
        write_cell(row.cells[1], value, font_size=9.5)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_row_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_table_geometry(table, widths_dxa: list[int], *, table_width_dxa: int) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(table_width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths_dxa):
            row.cells[idx].width = Inches(width / DXA_PER_INCH)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def write_cell(
    cell,
    text: str,
    *,
    bold: bool = False,
    font_size: float = 9.5,
    color: RGBColor = BLACK,
    align: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    parts = str(text).split("\n") if text else [""]
    for idx, part in enumerate(parts):
        paragraph = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        if align is not None:
            paragraph.alignment = align
        run = paragraph.add_run(part.strip())
        set_run_font(run, size=font_size, color=color, bold=bold)


def add_body_paragraph(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(_strip_markdown_emphasis(text))
    set_run_font(run, size=11, color=BLACK, bold=bold, italic=italic)


def add_quote(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [CONTENT_WIDTH_DXA], table_width_dxa=CONTENT_WIDTH_DXA)
    shade_cell(table.rows[0].cells[0], NOTE_FILL)
    write_cell(table.rows[0].cells[0], _strip_markdown_emphasis(text), font_size=9.5, color=MUTED)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(_strip_markdown_emphasis(text))
    set_run_font(run, size=11, color=BLACK)


def add_code_block(doc: Document, block: MarkdownBlock) -> None:
    label = "Diagram source" if block.language.lower() == "mermaid" else "Evidence / code block"
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(label)
    set_run_font(run, size=9, color=MUTED, bold=True)

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [CONTENT_WIDTH_DXA], table_width_dxa=CONTENT_WIDTH_DXA)
    shade_cell(table.rows[0].cells[0], LIGHT_FILL)
    write_cell(table.rows[0].cells[0], block.text, font_size=8.5)
    for paragraph in table.rows[0].cells[0].paragraphs:
        for run in paragraph.runs:
            set_run_font(run, size=8.5, color=BLACK, font_name="Consolas")


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    column_count = len(rows[0])
    table_width = CONTENT_WIDTH_DXA
    widths = _estimate_column_widths(rows, table_width)
    font_size = 7.8 if column_count >= 7 else 8.8 if column_count >= 5 else 9.5

    table = doc.add_table(rows=1, cols=column_count)
    table.style = "Table Grid"
    set_table_geometry(table, widths, table_width_dxa=table_width)
    set_row_repeat_header(table.rows[0])
    for idx, header in enumerate(rows[0]):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, HEADER_FILL)
        write_cell(cell, _strip_markdown_emphasis(header), bold=True, font_size=font_size, align=_alignment_for_column(header))

    for values in rows[1:]:
        row = table.add_row()
        for idx, value in enumerate(values):
            write_cell(row.cells[idx], _strip_markdown_emphasis(value), font_size=font_size, align=_alignment_for_column(rows[0][idx]))
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def _estimate_column_widths(rows: list[list[str]], total_dxa: int) -> list[int]:
    column_count = len(rows[0])
    scores: list[int] = []
    for col in range(column_count):
        values = [row[col] for row in rows if col < len(row)]
        max_len = max([len(value) for value in values] + [8])
        header_len = len(rows[0][col])
        score = min(max(max_len, header_len * 2, 8), 55)
        if rows[0][col].lower() in {"id", "no.", "step", "version", "date", "status", "risk"}:
            score = min(score, 16)
        scores.append(score)
    total_score = sum(scores)
    raw = [max(720, int(total_dxa * score / total_score)) for score in scores]
    delta = total_dxa - sum(raw)
    raw[-1] += delta
    return raw


def _alignment_for_column(header: str):
    normalized = header.strip().lower()
    if normalized in {"id", "no.", "step", "version", "date", "status", "risk", "answer", "result", "coverage"}:
        return WD_ALIGN_PARAGRAPH.CENTER
    if "date" in normalized or "status" in normalized or normalized in {"s/p/d"}:
        return WD_ALIGN_PARAGRAPH.CENTER
    return WD_ALIGN_PARAGRAPH.LEFT


def _strip_markdown_emphasis(text: str) -> str:
    return text.replace("**", "").replace("__", "").strip()


def render_blocks(doc: Document, blocks: list[MarkdownBlock]) -> None:
    first_h1_skipped = False
    for block in blocks:
        if block.kind == "heading":
            if block.level == 1 and not first_h1_skipped:
                first_h1_skipped = True
                continue
            level = min(max(block.level - 1, 1), 3)
            doc.add_heading(block.text, level=level)
        elif block.kind == "paragraph":
            add_body_paragraph(doc, block.text)
        elif block.kind == "quote":
            add_quote(doc, block.text)
        elif block.kind == "bullet":
            add_bullet(doc, block.text)
        elif block.kind == "table":
            add_markdown_table(doc, block.rows)
        elif block.kind == "code":
            add_code_block(doc, block)


def convert_one(source: Path, output_dir: Path, *, fix_known_issues: bool = False) -> Path:
    text = source.read_text(encoding="utf-8-sig")
    if fix_known_issues:
        text = apply_known_markdown_repairs(source, text)
    meta, body = parse_front_matter(text)
    validate_metadata(meta, source.name)
    blocks = parse_markdown_blocks(body, source_name=source.name)

    doc = Document()
    configure_document(doc)
    set_core_properties(doc, meta)
    add_header_footer(doc, meta)
    add_title_block(doc, meta)
    render_blocks(doc, blocks)

    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / source.with_suffix(".docx").name
    doc.save(output_path)
    return output_path


def convert_many(source_dir: Path, output_dir: Path, *, fix_known_issues: bool = False) -> list[Path]:
    if not source_dir.exists():
        raise FileNotFoundError(f"Source directory not found: {source_dir}")
    sources = sorted(source_dir.glob("*.md"))
    if not sources:
        raise FileNotFoundError(f"No Markdown files found in {source_dir}")
    return [convert_one(source, output_dir, fix_known_issues=fix_known_issues) for source in sources]


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Convert CSV108 Markdown files to pharmaceutical-style DOCX files.")
    parser.add_argument("--source-dir", type=Path, default=Path("CSV108"), help="Directory containing Markdown files.")
    parser.add_argument("--source", type=Path, help="Optional single Markdown file to convert.")
    parser.add_argument("--output-dir", type=Path, default=Path("CSV108_docx"), help="Directory for generated DOCX files.")
    parser.add_argument(
        "--fix-known-markdown-issues",
        action="store_true",
        help="Apply safe repairs for known CSV108 Markdown table issues before conversion.",
    )
    return parser


def main(argv: Iterable[str] | None = None) -> int:
    parser = build_arg_parser()
    args = parser.parse_args(list(argv) if argv is not None else None)

    try:
        if args.source:
            if not args.source.exists():
                raise FileNotFoundError(f"Source file not found: {args.source}")
            outputs = [convert_one(args.source, args.output_dir, fix_known_issues=args.fix_known_markdown_issues)]
        else:
            outputs = convert_many(args.source_dir, args.output_dir, fix_known_issues=args.fix_known_markdown_issues)
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1

    print(f"Generated {len(outputs)} DOCX file{'s' if len(outputs) != 1 else ''} in {args.output_dir}")
    for path in outputs[:8]:
        print(path)
    if len(outputs) > 8:
        print(f"... {len(outputs) - 8} more")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
